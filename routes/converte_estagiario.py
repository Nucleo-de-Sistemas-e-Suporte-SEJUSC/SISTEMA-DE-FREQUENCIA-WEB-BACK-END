#from utils.valida_ambiente_inux import valida_ambiente_pdf_linux
from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
from mysql.connector import Error
from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import zipfile 
import datetime
from datetime import time, timedelta, datetime
from datetime import date
from dateutil.easter import easter
import holidays
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


bp_converte_estagiario_pdf = Blueprint('bp_converte_estagiario_pdf', __name__)



def set_cell_background(cell, color_hex):
    """
    Define a cor de fundo da célula (color_hex no formato 'RRGGBB', ex: 'B7DEE8' para azul claro).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)

def pegar_feriados_mes(ano, mes, estado='AM'):
    print(f"DEBUG: Iniciando pegar_feriados_mes para ano={ano}, mes={mes}, estado='{estado}'") # DEBUG

    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = [] # Inicializa para o caso de falha na query
    try:
        query_sql = "SELECT data FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        print(f"DEBUG: Executando SQL: {query_sql} com params {params}") # DEBUG
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
        print(f"DEBUG: Feriados municipais crus do DB: {feriados_municipais_db}") # DEBUG
        if feriados_municipais_db:
            # Itera sobre uma cópia ou acessa diretamente, mas vamos ver o tipo do primeiro, se existir
            print(f"DEBUG: Tipo do valor 'data' do primeiro feriado do DB (se existir): {type(feriados_municipais_db[0]['data'])}") # DEBUG
    except Exception as e:
        print(f"DEBUG: Erro ao buscar feriados municipais do DB: {e}") #DEBUG
    finally:
        if conexao.is_connected(): # Verifica se a conexão está aberta antes de fechar
            cursor.close() # Fecha o cursor primeiro
            conexao.close()
            print("DEBUG: Conexão com MySQL fechada.") # DEBUG
        else:
            print("DEBUG: Conexão com MySQL já estava fechada ou não foi estabelecida.") #DEBUG


    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        print(f"DEBUG: Processando feriado_row['data']: {data_db} (Tipo: {type(data_db)})") # DEBUG
        
        data_feriado_obj = None
        if data_db is None:
            print(f"DEBUG: data_db é None. Pulando.") #DEBUG
            continue

        if hasattr(data_db, 'date'):  # Verifica se é um objeto datetime.datetime
            data_feriado_obj = data_db.date()
            print(f"DEBUG: Convertido de datetime.datetime para date: {data_feriado_obj}") # DEBUG
        elif isinstance(data_db, date):  # Verifica se já é um objeto datetime.date
            data_feriado_obj = data_db
            print(f"DEBUG: Já é um objeto date: {data_feriado_obj}") # DEBUG
        else:
            # Caso seja uma string ou outro tipo, tenta converter
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
                print(f"DEBUG: Convertido de string/outro para date: {data_feriado_obj}") # DEBUG
            except ValueError:
                print(f"DEBUG: Alerta: Formato de data inválido '{data_db}' não pôde ser convertido.") # DEBUG
                continue # Pula para o próximo feriado

        if data_feriado_obj:
            br_feriados[data_feriado_obj] = "Feriado Municipal"
            print(f"DEBUG: Adicionado ao br_feriados: {data_feriado_obj}") # DEBUG

    print(f"DEBUG: Conteúdo de br_feriados ANTES de filtrar por mês: {br_feriados.items()}") # DEBUG
    feriados_mes = [d for d in br_feriados if d.month == mes]
    print(f"DEBUG: Feriados filtrados para o mês {mes}: {feriados_mes}") # DEBUG
    return feriados_mes



def formatar_horario_para_hh_mm_v2(valor_horario):
    """
    Formata um valor de horário para o formato HH:MM, removendo os segundos.
    """
    if not valor_horario:  # Se for None, string vazia, etc.
        return ''

    # Caso 1: Se for um objeto datetime.time
    if isinstance(valor_horario, time):
        return valor_horario.strftime('%H:%M')

    # Caso 2: Se for um objeto datetime.timedelta (comum de bancos de dados para colunas TIME)
    if isinstance(valor_horario, timedelta):
        total_seconds = int(valor_horario.total_seconds())
        # Ignora dias, foca apenas na parte de tempo do dia
        if total_seconds < 0: # Lida com timedeltas negativos, se aplicável
            # Você pode querer um tratamento específico aqui, por ex., '' ou erro
            # Para simplificar, vamos assumir horas e minutos a partir de 0 se for negativo
            # ou tratar como 00:00. A lógica exata pode depender do seu caso de uso.
            # Exemplo: tratar como 00:00 se negativo ou converter para positivo
            # Para este exemplo, vamos apenas calcular com base no valor absoluto.
            total_seconds = abs(total_seconds)

        hours = (total_seconds // 3600) % 24 # Garante que as horas fiquem dentro de 0-23
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"

    # Caso 3: Se for uma string
    if isinstance(valor_horario, str):
        try:
            # Tenta primeiro como HH:MM:SS
            if valor_horario.count(':') == 2:
                dt_obj = datetime.strptime(valor_horario, '%H:%M:%S')
                return dt_obj.strftime('%H:%M')
            # Depois como HH:MM
            elif valor_horario.count(':') == 1:
                dt_obj = datetime.strptime(valor_horario, '%H:%M')
                return dt_obj.strftime('%H:%M') # Já está no formato, mas re-formata para garantir
            else:
                # Se não for um formato de tempo reconhecido, retorna a string original
                return valor_horario
        except ValueError:
            # Se a conversão da string falhar
            return valor_horario # Retorna a string original

    # Fallback: Se não for nenhum dos tipos acima, tenta converter para string
    return str(valor_horario)

@bp_converte_estagiario_pdf.route('/api/estagiario/pdf', methods=['POST'])
def converte_estagiario_pdf():
    try:
        body = request.json or {}
        estagiarios_id = body.get('estagiarios', [])
        print(estagiarios_id)

        if not estagiarios_id:
            return jsonify({'erro': 'Nenhum estagiário selecionado'}), 400

        try:
            ids = [int(id) for id in estagiarios_id]
        except ValueError:
            return jsonify({'erro': 'IDs inválidos'}), 400

        mes_body = body.get('mes')
        data_ano_mes_atual = data_atual(mes_body)
        mes_por_extenso = data_ano_mes_atual['mes']
        mes_numerico = data_ano_mes_atual['mes_numerico']
        ano = data_ano_mes_atual['ano']

        conexao = connect_mysql()
        cursor = conexao.cursor(dictionary=True)

        placeholders = ','.join(['%s'] * len(ids))
        query = f"SELECT * FROM estagiarios WHERE id IN ({placeholders})"
        cursor.execute(query, ids)
        estagiarios = cursor.fetchall()


        if not estagiarios:
            conexao.close()
            return jsonify({'erro': 'Nenhum estagiário encontrado'}), 404

        arquivos_gerados = []
        
        estado_para_feriados = 'AM' # Ou defina dinamicamente se necessário

        feriados_mes_corrente_periodo = pegar_feriados_mes(ano, mes_numerico, estado=estado_para_feriados) # Feriados do primeiro mês do período (ex: Junho)

        ano_proximo_mes_periodo = ano
        mes_numerico_proximo_periodo = mes_numerico + 1
        if mes_numerico_proximo_periodo > 12:
            mes_numerico_proximo_periodo = 1
            ano_proximo_mes_periodo += 1
        
        feriados_proximo_mes_periodo = pegar_feriados_mes(ano_proximo_mes_periodo, mes_numerico_proximo_periodo, estado=estado_para_feriados) # Feriados do segundo mês do período (ex: Julho)
        
        # Combina as duas listas de feriados. Usar set para evitar duplicatas caso haja alguma sobreposição (improvável com meses distintos)
        todos_feriados_do_periodo = list(set(feriados_mes_corrente_periodo + feriados_proximo_mes_periodo))
        print(f"DEBUG: Todos feriados para o período (mes {mes_numerico} e {mes_numerico_proximo_periodo}): {todos_feriados_do_periodo}")
        # -- Fim da Modificação --
        #feriados_do_mes = pegar_feriados_mes(ano, mes_numerico)

        for estagiario in estagiarios:
            template_path = 'FREQUÊNCIA ESTAGIÁRIOS - MODELO.docx'
            doc = Document(template_path)

            cria_dias_da_celula(doc, ano, mes_numerico, estagiario, todos_feriados_do_periodo)

            troca_de_dados = {
                "CAMPO SETOR": estagiario['setor'],
                "CAMPO MES": mes_por_extenso,
                "CAMPO NOME": estagiario['nome'],
                #"CAMPO PERIODO": f"21/{mes_numerico}/{ano} a 20/{(mes_numerico % 12) + 1}/{ano if mes_numerico < 12 else ano + 1}",
                "CAMPO ANO": str(ano),
                "CAMPO HORARIO": str(estagiario.get('horario')),
                "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_entrada')),
                "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(estagiario.get('horario_saida')),
                "CAMPO CARGO": str(estagiario.get('cargo')),
            }

            for placeholder, valor in troca_de_dados.items():
                muda_texto_documento(doc, placeholder, valor)
            nome_limpo = estagiario['nome'].strip()
            setor_limpo = estagiario['setor'].replace('/', '-').strip()
            caminho_pasta = f"setor/{setor_limpo}/estagiario/{mes_por_extenso}/{nome_limpo}"
            os.makedirs(caminho_pasta, exist_ok=True)

            nome_base = f"{nome_limpo.replace(' ', '_')}_FREQUENCIA"
            docx_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.docx"))
            pdf_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.pdf"))

            
            doc.save(docx_path)
            convert_to_pdf(docx_path, pdf_path)

            arquivos_gerados.append(pdf_path)
            


        # Cria um arquivo ZIP com todos os PDFs
        zip_path = f"setor/frequencias_{mes_por_extenso}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for pdf in arquivos_gerados:
                zipf.write(pdf, os.path.basename(pdf))

        # Salva o ZIP no banco
        cursor.execute(
            "INSERT INTO arquivos_zip (mes, caminho_zip, tipo) VALUES (%s, %s,%s)",
            (mes_por_extenso, zip_path, 'estagiario')
        )

        conexao.commit()
        conexao.close()

        return send_file(
            zip_path,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'frequencias_estagiarios_{mes_por_extenso}.zip'
        )

    except Exception as exception:
        if 'conexao' in locals():
            conexao.close()
        return jsonify({'erro': f'Erro: {str(exception)}'}), 500
    
    
def cria_dias_da_celula(doc, ano, mes_numerico, estagiario, feriados):
    from datetime import datetime, timedelta, date
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    def calcula_periodo_21_a_20(ano_calc, mes_calc):
        data_inicio = datetime(ano_calc, mes_calc, 21)
        if mes_calc == 12:
            data_fim = datetime(ano_calc + 1, 1, 20)
        else:
            data_fim = datetime(ano_calc, mes_calc + 1, 20)

        dias_periodo = []
        data_iter_calc = data_inicio
        while data_iter_calc <= data_fim:
            dias_periodo.append({
                "dia": data_iter_calc.day,
                "mes": data_iter_calc.month,
                "ano": data_iter_calc.year
            })
            data_iter_calc += timedelta(days=1)
        return dias_periodo

    linha_inicial = 7
    table = doc.tables[0] # Pega a primeira tabela do documento

    # Configuração inicial das linhas existentes
    for row in table.rows:
        row.height = Cm(0.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7)
                    run.font.bold = False

    dias_periodo = calcula_periodo_21_a_20(ano, mes_numerico)

    # ### LÓGICA DE CRIAÇÃO DE LINHAS REMOVIDA DAQUI ###
    # O código que adicionava novas linhas foi apagado, conforme solicitado.

    # Loop principal para preencher os dados de cada dia do período
    for i, dia_info in enumerate(dias_periodo):
        # ... (O restante do seu loop que preenche os dias continua aqui, sem alterações)
        dia = dia_info["dia"]
        mes_iter = dia_info["mes"]
        ano_dia = dia_info["ano"]
        
        data_iteracao_atual = date(ano_dia, mes_iter, dia)
        # Assumindo que a função pega_final_de_semana existe em outro lugar
        dia_semana = pega_final_de_semana(ano_dia, mes_iter, dia)
        
        row = table.rows[linha_inicial + i]
        
        for cell in row.cells:
            cell.text = ""
            for paragraph in cell.paragraphs:
                paragraph.clear()
        
        dia_cell = row.cells[0]
        dia_paragraph = dia_cell.paragraphs[0]
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        is_feriado = data_iteracao_atual in feriados
        is_ferias = False
        if estagiario.get('feriasinicio') and estagiario.get('feriasfinal'):
            ferias_inicio = estagiario['feriasinicio']
            if isinstance(ferias_inicio, datetime):
                ferias_inicio = ferias_inicio.date()
            
            ferias_final = estagiario['feriasfinal']
            if isinstance(ferias_final, datetime):
                ferias_final = ferias_final.date()

            if ferias_inicio <= data_iteracao_atual <= ferias_final:
                is_ferias = True

        texto = None
        if dia_semana == 5: texto = "SÁBADO"
        elif dia_semana == 6: texto = "DOMINGO"
        elif is_ferias and dia_semana not in [5, 6]: texto = "FÉRIAS"
        elif is_feriado and dia_semana not in [5, 6]: texto = "FERIADO"

        if texto:
            # Assumindo que a função set_row_background existe
            set_row_background(row, 'C5E0B4')
            celulas_para_marcar = [2, 5, 8, 12]
            if texto == "FERIADO":
                 celulas_para_marcar = [2, 5, 9]

            for j in celulas_para_marcar:
                if j < len(row.cells):
                    cell = row.cells[j]
                    for p in cell.paragraphs:
                        p.clear()
                    p_cell = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    run_cell = p_cell.add_run(texto)
                    run_cell.font.bold = True
                    run_cell.font.name = "Calibri"
                    run_cell.font.size = Pt(7)
                    p_cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # ... (Fim do loop)

    # ### NOVA LÓGICA PARA REMOVER LINHAS EXTRAS ###
    # Este bloco foi adicionado conforme solicitado.
    
    # Calcula o número de linhas de dados que o template possui
    total_linhas_dados_template = len(table.rows) - linha_inicial
    
    # Calcula o número de dias (e, portanto, linhas) que o período atual realmente usou
    dias_no_periodo_atual = len(dias_periodo)

    # Se o número de linhas no template for maior que o necessário, remove as que sobraram
    if total_linhas_dados_template > dias_no_periodo_atual:
        linhas_para_remover = total_linhas_dados_template - dias_no_periodo_atual
        
        for _ in range(linhas_para_remover):
            # Pega a referência do elemento da última linha (tr) e remove da tabela (tbl)
            ultima_linha = table.rows[-1]
            tr_element = ultima_linha._tr
            tbl_element = table._tbl
            tbl_element.remove(tr_element)
