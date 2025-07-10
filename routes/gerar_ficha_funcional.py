import os
import uuid
from flask import Blueprint, jsonify, request
from conection_mysql import connect_mysql
from docx import Document
from utils.gerador_docx import preencher_documento

bp_gerar_ficha_funcional = Blueprint("bp_gerar_ficha_funcional", __name__)
UPLOADS_FOLDER = os.path.join(os.getcwd(), "uploads")

@bp_gerar_ficha_funcional.route("/api/servidores/<int:funcionario_id>/gerar-ficha-funcional", methods=["POST"])
def gerar_ficha(funcionario_id):
    try:
        conexao = connect_mysql()
        cursor = conexao.cursor(dictionary=True)
        cursor.execute("SELECT * FROM funcionarios WHERE id = %s", (funcionario_id,))
        funcionario = cursor.fetchone()

        if not funcionario:
            return jsonify({"erro": "Funcionário não encontrado"}), 404

        template_path = "FICHA FUNCIONAL (1).docx"
        if not os.path.exists(template_path):
            return jsonify({"erro": "Template 'FICHA FUNCIONAL (1).docx' não encontrado."}), 500
        
        doc = Document(template_path)

        # ------------------- INÍCIO DO CÓDIGO DE DIAGNÓSTICO -------------------
        print("\n--- INICIANDO DIAGNÓSTICO DO DOCUMENTO ---")
        print("--- PARÁGRAFOS PRINCIPAIS ---")
        for i, p in enumerate(doc.paragraphs):
            print(f"P{i}: '{p.text}'")

        print("\n--- TEXTO DAS CÉLULAS DAS TABELAS ---")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        print(f"T{t_idx}-R{r_idx}-C{c_idx}-P{p_idx}: '{p.text}'")
        print("--- FIM DO DIAGNÓSTICO ---\n")
        # -------------------- FIM DO CÓDIGO DE DIAGNÓSTICO --------------------

        mapeamento = {
            "CAMPO_NOME": funcionario.get("nome"),
            "CAMPO_DATA_NASCIMENTO": str(funcionario.get("data_nascimento", "")),
            "CAMPO_SEXO": funcionario.get("sexo"),
            "CAMPO_ESTADO_CIVIL": funcionario.get("estado_civil"),
            "CAMPO_NATURALIDADE": funcionario.get("naturalidade"),
            "CAMPO_NACIONALidade": funcionario.get("nacionalidade"), # <-- Atenção aqui
            "CAMPO_CPF": funcionario.get("cpf"),
            "CAMPO_PIS": funcionario.get("pis"),
            "CAMPO_IDENTIDADE": funcionario.get("identidade"),
            "CAMPO_TITULO": funcionario.get("titulo_eleitor"),
            "CAMPO_DATA_ADMISSAO": str(funcionario.get("data_Admissao", "")),
            "CAMPO_CARGO": funcionario.get("cargo"),
            "CAMPO_ENDERECO": funcionario.get("endereco"),
            "CAMPO_PAI": funcionario.get("nome_pai"),
            "CAMPO_MAE": funcionario.get("nome_mae"),
            "CAMPO_SERVICO_MILITAR": funcionario.get("servico_militar"),
            "CAMPO_CARTEIRA_PROF": funcionario.get("carteira_profissional"),
            "CAMPO_DATA_POSSE": str(funcionario.get("data_posse", "")),
        }
        
        preencher_documento(doc, mapeamento)

        # O resto do código continua igual...
        nome_original = f"Ficha_Funcional_{funcionario['nome'].replace(' ', '_')}.docx"
        nome_unico = f"{uuid.uuid4()}.docx"
        caminho_salvo = os.path.join(UPLOADS_FOLDER, nome_unico)
        doc.save(caminho_salvo)

        query_insert = """
            INSERT INTO documentos 
            (nome_original, nome_armazenado, caminho_arquivo, tipo_documento, funcionario_id)
            VALUES (%s, %s, %s, %s, %s)
        """
        cursor.execute(query_insert, (nome_original, nome_unico, caminho_salvo, "Ficha Funcional", funcionario_id))
        conexao.commit()
        documento_id = cursor.lastrowid
        
        cursor.close()
        conexao.close()

        return jsonify({
            "mensagem": "Ficha Funcional gerada com sucesso!",
            "documento_id": documento_id,
            "caminho": caminho_salvo
        }), 201

    except Exception as e:
        return jsonify({"erro": f"Ocorreu um erro: {str(e)}"}), 500