from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
from mysql.connector import Error
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import holidays
from datetime import datetime, date, timedelta,time
import zipfile
import re
from datetime import date
from dateutil.easter import easter
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


bp_converte_servidor_pdf = Blueprint('bp_converte_servidor_pdf', __name__)



def set_cell_background(cell, color_hex):
  
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_background(row, color_hex):
    for cell in row.cells:
        set_cell_background(cell, color_hex)
    
# ...existing code...
def pegar_feriados_mes(ano, mes, estado='AM'):
    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao = connect_mysql()
    cursor = conexao.cursor(dictionary=True)
    feriados_municipais_db = []
    try:
        query_sql = "SELECT data, ponto_facultativo FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        cursor.execute(query_sql, params)
        feriados_municipais_db = cursor.fetchall()
    finally:
        if conexao.is_connected():
            cursor.close()
            conexao.close()

    pontos_facultativos = []
    feriados_normais = []
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        ponto_facultativo = feriado_row.get('ponto_facultativo', 0)
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                continue
        if data_feriado_obj:
            if ponto_facultativo:
                pontos_facultativos.append(data_feriado_obj)
            else:
                feriados_normais.append(data_feriado_obj)
            br_feriados[data_feriado_obj] = "Feriado Municipal"
    feriados_mes = [d for d in br_feriados if d.month == mes]
    pontos_facultativos_mes = [d for d in pontos_facultativos if d.month == mes]
    return feriados_mes, pontos_facultativos_mes

def limpa_nome(nome):
    return re.sub(r'[^\w\s-]', '', nome).strip().replace(' ', '_')


def formatar_horario_para_hh_mm_v2(valor_horario):
 
    if not valor_horario: 
        return ''

   
    if isinstance(valor_horario, time):
        return valor_horario.strftime('%H:%M')

   
    if isinstance(valor_horario, timedelta):
        total_seconds = int(valor_horario.total_seconds())
       
        if total_seconds < 0: 
        
            total_seconds = abs(total_seconds)

        hours = (total_seconds // 3600) % 24 # Garante que as horas fiquem dentro de 0-23
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"

 
    if isinstance(valor_horario, str):
        try:
           
            if valor_horario.count(':') == 2:
                dt_obj = datetime.strptime(valor_horario, '%H:%M:%S')
                return dt_obj.strftime('%H:%M')
           
            elif valor_horario.count(':') == 1:
                dt_obj = datetime.strptime(valor_horario, '%H:%M')
                return dt_obj.strftime('%H:%M') 
            else:
               
                return valor_horario
        except ValueError:
          
            return valor_horario 

   
    return str(valor_horario)

@bp_converte_servidor_pdf.route('/api/servidores/pdf', methods=['POST'])
def converte_servidor_pdf():
    try:
        body = request.json or {}
        funcionarios_id = body.get('funcionarios', [])

        if not funcionarios_id:
            return jsonify({'erro': 'Nenhum funcionário selecionado'}), 400

        try:
            ids = [int(id) for id in funcionarios_id]
        except ValueError:
            return jsonify({'erro': 'IDs inválidos'}), 400

        mes_body = body.get('mes')
        data_ano_mes_atual = data_atual(mes_body)
        mes_por_extenso = data_ano_mes_atual['mes']
        mes_numerico = data_ano_mes_atual['mes_numerico']
        ano = data_ano_mes_atual['ano']
        quantidade_dias_no_mes = pega_quantidade_dias_mes(ano, mes_numerico)

        conexao = connect_mysql()
        cursor = conexao.cursor(dictionary=True)

        placeholders = ','.join(['%s'] * len(ids))
        query = f"SELECT * FROM funcionarios WHERE id IN ({placeholders})"
        cursor.execute(query, ids)
        funcionarios = cursor.fetchall()

        if not funcionarios:
            conexao.close()
            print(" cair aquiiii")
            return jsonify({'erro': 'Nenhum funcionário encontrado'}), 404

        arquivos_gerados = []
    
        for funcionario in funcionarios:
            estado_funcionarios = funcionario.get('estado', 'AM')
            feriados_do_mes, pontos_facultativos_mes = pegar_feriados_mes(ano, mes_numerico, estado=estado_funcionarios)
            template_path = 'FREQUÊNCIA_MENSAL.docx'
            doc = Document(template_path)

            cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados_do_mes, pontos_facultativos_mes)
          
            troca_de_dados = {
            "CAMPO SETOR": funcionario['setor'],
            "CAMPO MÊS": mes_por_extenso,
            "CAMPO NOME": funcionario['nome'],
            "CAMPO ANO": str(ano),
            "CAMPO HORARIO": funcionario.get('horario', ''), # Mantido como está, sem formatação específica aqui
            "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(funcionario.get('horarioentrada', '')),
            "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(funcionario.get('horariosaida', '')),
            "CAMPO MATRÍCULA": str(funcionario.get('matricula', '')),
            "CAMPO CARGO": funcionario.get('cargo', ''),
        }

            for placeholder, valor in troca_de_dados.items():
                muda_texto_documento(doc, placeholder, valor)

            nome_limpo = limpa_nome(funcionario['nome'])
            setor_limpo = limpa_nome(funcionario['setor'])
            caminho_pasta = f"setor/{setor_limpo}/servidor/{mes_por_extenso}/{nome_limpo}"
            os.makedirs(caminho_pasta, exist_ok=True)

            nome_base = f"{nome_limpo}_FREQUENCIA"
            docx_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.docx"))
            pdf_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.pdf"))

            doc.save(docx_path)
            convert_to_pdf(docx_path, pdf_path)

            arquivos_gerados.append(pdf_path)

     
        zip_path = os.path.abspath(f"setor/frequencias_{mes_por_extenso}.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for pdf in arquivos_gerados:
                zipf.write(pdf, os.path.basename(pdf))

 
        cursor.execute(
            "INSERT INTO arquivos_zip (mes, caminho_zip, tipo) VALUES (%s, %s, %s)",
            (mes_por_extenso, zip_path, 'servidores')
        )

        conexao.commit()
        conexao.close()

        return send_file(
            zip_path,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'frequencias_servidores_{mes_por_extenso}.zip'
        )

    except Exception as exception:
        if 'conexao' in locals():
            conexao.close()
        return jsonify({'erro': f'Erro: {str(exception)}'}), 500

def cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados,pontos_facultativos):
    linha_inicial = 8

    if not doc.tables:
        print("AVISO: Nenhum tabela encontrada no documento.")
        return
    
    table = doc.tables[0] 

    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7)
                    run.font.bold = False

    target_total_rows_in_table = linha_inicial + quantidade_dias_no_mes


    while len(table.rows) > target_total_rows_in_table:
        row_to_delete = table.rows[-1] 
        tbl_element = table._tbl
        tr_element = row_to_delete._tr
        tbl_element.remove(tr_element)
        print(f"INFO: Linha excedente removida. Total de linhas agora: {len(table.rows)}")

   
    while len(table.rows) < target_total_rows_in_table:
        new_row = table.add_row()
        new_row.height = Cm(0.5) 
        new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for cell in new_row.cells:

            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
       
        print(f"INFO: Linha faltante adicionada. Total de linhas agora: {len(table.rows)}")



    for i in range(quantidade_dias_no_mes):
        dia = i + 1
     
        row = table.rows[linha_inicial + i]
        data_atual = date(ano, mes_numerico, dia) 
        dia_semana = pega_final_de_semana(ano, mes_numerico, dia) 
        # Limpeza das células da linha atual antes de preencher
        for cell in row.cells:
            cell.text = "" 
            for paragraph in cell.paragraphs: 
                paragraph.clear()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

     
        dia_cell = row.cells[0]
     
        dia_paragraph = dia_cell.paragraphs[0] if dia_cell.paragraphs else dia_cell.add_paragraph()
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
      
        
        texto_status = "" # Para Sábado/Domingo

        if dia_semana == 5:
            texto_status = "SÁBADO"
        elif dia_semana == 6:
            texto_status = "DOMINGO"

        if texto_status: 
            set_row_background(row, 'C5E0B4') # VERDE
            for j in [2, 5, 9, 13]: 
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = texto_status
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs: 
                            run.font.bold = True
                            run.font.name = "Calibri" 
                            run.font.size = Pt(7)    
                else:
                    print(f"AVISO: Índice de coluna {j} para S/D fora dos limites.")

     
        if data_atual in pontos_facultativos and dia_semana not in [5, 6]:
            set_row_background(row, 'C5E0B4')  # 
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "PONTO FACULTATIVO"
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(7)
                else:
                    print(f"AVISO: Índice de coluna {j} para PONTO FACULTATIVO fora dos limites.")


        elif data_atual in feriados and dia_semana not in [5, 6]:
            set_row_background(row, 'C5E0B4') # VERDE
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "FERIADO"

                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(7)
                else:
                    print(f"AVISO: Índice de coluna {j} para FERIADO fora dos limites.")


        # Férias (exceto fins de semana) - sobrescreve células se for o caso
        if funcionario.get('feriasinicio') and funcionario.get('feriasfinal'):
            ferias_inicio_raw = funcionario['feriasinicio']
            ferias_final_raw = funcionario['feriasfinal']
            ferias_inicio = ferias_inicio_raw.date() if hasattr(ferias_inicio_raw, 'date') else ferias_inicio_raw
            ferias_final = ferias_final_raw.date() if hasattr(ferias_final_raw, 'date') else ferias_final_raw

            if isinstance(ferias_inicio, date) and isinstance(ferias_final, date) and \
               (ferias_inicio <= data_atual <= ferias_final and dia_semana not in [5, 6]):
                set_row_background(row, 'C5E0B4') # VERDE   
                for j in [2, 5, 9, 13]:
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = "FÉRIAS"
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = "Calibri"
                                run.font.size = Pt(7)
                    else:
                        print(f"AVISO: Índice de coluna {j} para FÉRIAS fora dos limites.")
