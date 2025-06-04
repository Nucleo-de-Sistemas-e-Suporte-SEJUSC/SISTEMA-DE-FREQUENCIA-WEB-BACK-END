from utils.convert_to_pdf import convert_to_pdf
from utils.muda_texto_documento import muda_texto_documento
from utils.formata_datas import data_atual, pega_final_de_semana, pega_quantidade_dias_mes
from flask import Blueprint, request, jsonify, send_file
from conection_mysql import connect_mysql
# from mysql.connector import Error # Não usado diretamente
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import zipfile
from datetime import datetime, date, time, timedelta
# Imports necessários para pegar_feriados_mes
from dateutil.easter import easter
import holidays
# import uuid # uuid não está sendo usado

bp_converte_setor_pdf = Blueprint('bp_converte_setor_pdf', __name__)

def pegar_feriados_mes(ano, mes, estado='AM'):
    print(f"DEBUG setor_funcionarios: Iniciando pegar_feriados_mes para ano={ano}, mes={mes}, estado='{estado}'")

    br_feriados = holidays.Brazil(state=estado)
    pascoa = easter(ano)
    corpus_christi = pascoa + timedelta(days=60)
    br_feriados[corpus_christi] = "Corpus Christi"

    conexao_feriado = None
    try:
        conexao_feriado = connect_mysql()
        cursor_feriado = conexao_feriado.cursor(dictionary=True)
        feriados_municipais_db = []
        query_sql = "SELECT data FROM feriados_municipais WHERE estado = %s AND YEAR(data) = %s"
        params = (estado, ano)
        print(f"DEBUG setor_funcionarios: Executando SQL: {query_sql} com params {params}")
        cursor_feriado.execute(query_sql, params)
        feriados_municipais_db = cursor_feriado.fetchall()
        print(f"DEBUG setor_funcionarios: Feriados municipais crus do DB: {feriados_municipais_db}")
        if feriados_municipais_db and feriados_municipais_db[0].get('data') is not None:
            print(f"DEBUG setor_funcionarios: Tipo do valor 'data' do primeiro feriado do DB (se existir): {type(feriados_municipais_db[0]['data'])}")
        cursor_feriado.close()
    except Exception as e:
        print(f"DEBUG setor_funcionarios: Erro ao buscar feriados municipais do DB: {e}")
    finally:
        if conexao_feriado and conexao_feriado.is_connected():
            conexao_feriado.close()
            print("DEBUG setor_funcionarios: Conexão de feriado com MySQL fechada.")
        else:
            print("DEBUG setor_funcionarios: Conexão de feriado com MySQL já estava fechada ou não foi estabelecida.")
            
    for feriado_row in feriados_municipais_db:
        data_db = feriado_row['data']
        data_feriado_obj = None
        if data_db is None:
            continue
        if hasattr(data_db, 'date'):
            data_feriado_obj = data_db.date()
        elif isinstance(data_db, date):
            data_feriado_obj = data_db
        else:
            try:
                data_feriado_obj = date.fromisoformat(str(data_db))
            except ValueError:
                print(f"DEBUG setor_funcionarios: Alerta: Formato de data inválido '{data_db}' não pôde ser convertido.")
                continue
        if data_feriado_obj:
            br_feriados[data_feriado_obj] = "Feriado Municipal"
    # print(f"DEBUG setor_funcionarios: Conteúdo de br_feriados ANTES de filtrar por mês: {br_feriados.items()}")
    feriados_mes_filtrados = [d for d in br_feriados if d.month == mes] # Renomeada a variável de retorno para clareza
    print(f"DEBUG setor_funcionarios: Feriados filtrados para o mês {mes}: {feriados_mes_filtrados}")
    return feriados_mes_filtrados
# FIM DE: Função pegar_feriados_mes


def formatar_horario_para_hh_mm_v2(valor_horario):
    if not valor_horario: return ''
    if isinstance(valor_horario, time): return valor_horario.strftime('%H:%M')
    if isinstance(valor_horario, timedelta):
        total_seconds = abs(int(valor_horario.total_seconds()))
        hours = (total_seconds // 3600) % 24
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02}:{minutes:02}"
    if isinstance(valor_horario, str):
        try:
            if valor_horario.count(':') == 2: return datetime.strptime(valor_horario, '%H:%M:%S').strftime('%H:%M')
            elif valor_horario.count(':') == 1: return datetime.strptime(valor_horario, '%H:%M').strftime('%H:%M')
            return valor_horario
        except ValueError: return valor_horario
    return str(valor_horario)

@bp_converte_setor_pdf.route('/api/setores/pdf', methods=['POST'])
def converte_setores_pdf():
    conexao_principal = None # Para ser usado no try/finally da rota
    try:
        body = request.json or {}
        setores = body.get('setores')
        mes_body = body.get('mes')

        if not setores or not isinstance(setores, list):
            return jsonify({'erro': 'Nenhum setor selecionado ou formato inválido'}), 400

        arquivos_zip_gerados_todos_setores = [] # Renomeado para clareza
        mes_por_extenso_geral = data_atual(mes_body)['mes'] # Pegar uma vez se o mês é o mesmo para todos

        for setor_nome in setores:
            # Os dados de data/mês/ano são por setor aqui, caso mes_body pudesse variar (embora improvável)
            data_ano_mes_atual = data_atual(mes_body)
            mes_numerico = data_ano_mes_atual['mes_numerico']
            ano = data_ano_mes_atual['ano']
            quantidade_dias_no_mes = pega_quantidade_dias_mes(ano, mes_numerico) 

            conexao_principal = connect_mysql() # Conexão por setor
            cursor = conexao_principal.cursor(dictionary=True)

            query_funcionarios = "SELECT * FROM funcionarios WHERE setor = %s"
            cursor.execute(query_funcionarios, (setor_nome,))
            funcionarios = cursor.fetchall()

            if not funcionarios:
                if conexao_principal and conexao_principal.is_connected():
                    cursor.close()
                    conexao_principal.close()
                print(f"Nenhum funcionário encontrado no setor {setor_nome}, pulando.")
                continue

            arquivos_pdf_gerados_neste_setor = [] # Renomeado para clareza
            setor_limpo = setor_nome.strip().replace('/', '_')

            for funcionario in funcionarios:
                # --- Busca de feriados por funcionário (considerando estado do funcionário) ---
                estado_funcionario = funcionario.get('estado', 'AM') # Pega o estado do funcionário, default 'AM'
                feriados_do_mes_funcionario = pegar_feriados_mes(ano, mes_numerico, estado=estado_funcionario)
                print(f"DEBUG ROTA setor_funcionarios: Feriados para {funcionario.get('nome')} no estado {estado_funcionario}: {feriados_do_mes_funcionario}")
                # --- FIM da busca de feriados ---

                template_path = 'FREQUÊNCIA_MENSAL.docx'
                doc = Document(template_path)
                
                # Passa a lista de feriados específica do funcionário
                cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados_do_mes_funcionario)

                troca_de_dados = {
                    "CAMPO SETOR": funcionario.get('setor', ''),
                    "CAMPO MÊS": mes_por_extenso_geral, # Usar o geral
                    "CAMPO NOME": funcionario.get('nome', ''),
                    "CAMPO ANO": str(ano),
                    "CAMPO HORARIO": str(funcionario.get('horario', '')),
                    "CAMPO ENTRADA": formatar_horario_para_hh_mm_v2(funcionario.get('horarioentrada', '')),
                    "CAMPO SAÍDA": formatar_horario_para_hh_mm_v2(funcionario.get('horariosaida', '')),
                    "CAMPO MATRÍCULA": str(funcionario.get('matricula', '')),
                    "CAMPO CARGO": funcionario.get('cargo', ''),
                }
                for placeholder, valor in troca_de_dados.items():
                    muda_texto_documento(doc, placeholder, valor)

                nome_limpo = funcionario.get('nome', 'NOME_PADRAO').strip().replace('/', '_')
                caminho_pasta = f"setor/{setor_limpo}/{mes_por_extenso_geral}"
                os.makedirs(caminho_pasta, exist_ok=True)
                nome_base = f"FREQUENCIA_{nome_limpo.replace(' ', '_')}"
                docx_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.docx"))
                pdf_path = os.path.abspath(os.path.join(caminho_pasta, f"{nome_base}.pdf"))
                doc.save(docx_path)
                convert_to_pdf(docx_path, pdf_path)
                arquivos_pdf_gerados_neste_setor.append(pdf_path)
                cursor.execute(
                    "INSERT INTO arquivos_pdf (servidor_id, caminho_pdf) VALUES (%s, %s)",
                    (funcionario['id'], pdf_path)
                )

            if arquivos_pdf_gerados_neste_setor: # Só cria ZIP se houver PDFs
                zip_path_setor = f"setor/{setor_limpo}/frequencias_funcionarios_{setor_limpo}_{mes_por_extenso_geral}.zip" # Nome mais específico
                with zipfile.ZipFile(zip_path_setor, 'w') as zipf:
                    for pdf in arquivos_pdf_gerados_neste_setor:
                        zipf.write(pdf, os.path.basename(pdf))
                arquivos_zip_gerados_todos_setores.append(zip_path_setor)
                cursor.execute(
                    "INSERT INTO arquivos_zip (setor, mes, caminho_zip, tipo) VALUES (%s, %s, %s, %s)",
                    (setor_nome, mes_por_extenso_geral, zip_path_setor, 'funcionarios_setor') # tipo ajustado
                )
            
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()

        if not arquivos_zip_gerados_todos_setores:
            return jsonify({'message': 'Nenhum arquivo ZIP de funcionários foi gerado (sem funcionários nos setores ou sem PDFs).'}), 200

        if len(arquivos_zip_gerados_todos_setores) > 1:
            zip_final_path = f"setor/frequencias_multissetores_funcionarios_{mes_body.replace('/','-')}_{ano}.zip" # Nome com ano e tipo
            with zipfile.ZipFile(zip_final_path, 'w') as zipf:
                for zip_file in arquivos_zip_gerados_todos_setores:
                    zipf.write(zip_file, os.path.basename(zip_file))
            
            conexao_principal = connect_mysql() # Nova conexão para salvar o ZIP agregado
            cursor = conexao_principal.cursor(dictionary=True)
            cursor.execute(
                "INSERT INTO arquivos_zip (setor, mes, ano, caminho_zip, tipo) VALUES (%s, %s, %s, %s, %s)", # Adicionado ano
                ('multissetores_funcionarios', mes_por_extenso_geral, str(ano), zip_final_path, 'multissetores_funcionarios_geral') # tipo ajustado
            )
            conexao_principal.commit()
            if conexao_principal and conexao_principal.is_connected():
                cursor.close()
                conexao_principal.close()
            return send_file(zip_final_path, mimetype='application/zip', as_attachment=True, download_name=os.path.basename(zip_final_path))
        elif arquivos_zip_gerados_todos_setores:
            return send_file(arquivos_zip_gerados_todos_setores[0], mimetype='application/zip', as_attachment=True, download_name=os.path.basename(arquivos_zip_gerados_todos_setores[0]))
        
        return jsonify({'message': 'Processamento concluído, mas nenhum ZIP para enviar.'}), 200

    except Exception as exception:
        print(f"ERRO ROTA SETOR FUNCIONARIOS: {str(exception)}")
        # import traceback; traceback.print_exc();
        if conexao_principal and conexao_principal.is_connected():
            cursor.close() # Garante que o cursor seja fechado se foi aberto
            conexao_principal.close()
        return jsonify({'erro': f'Erro ao processar setores: {str(exception)}'}), 500


# Modificado para aceitar 'feriados' e aplicar a lógica
def cria_dias_da_celula(doc, quantidade_dias_no_mes, ano, mes_numerico, funcionario, feriados):
    linha_inicial = 8

    if not doc.tables:
        print("AVISO: Nenhum tabela encontrada no documento.")
        return
    
    table = doc.tables[0] # Assume-se que a primeira tabela é a de frequência

    # 1. Aplicar formatação base em todas as linhas existentes (como no seu código original)
    # Esta formatação pode ser muito genérica; idealmente, o template já teria os estilos corretos
    # para cabeçalhos vs. dados, mas vamos manter sua lógica original por enquanto.
    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7)
                    run.font.bold = False

    # 2. Ajustar o número de linhas na tabela para corresponder à quantidade_dias_no_mes
    # Linhas de dados necessárias = quantidade_dias_no_mes
    # Total de linhas que a tabela deve ter = linha_inicial (cabeçalho) + quantidade_dias_no_mes
    target_total_rows_in_table = linha_inicial + quantidade_dias_no_mes

    # Remover linhas excedentes do final da tabela
    while len(table.rows) > target_total_rows_in_table:
        row_to_delete = table.rows[-1] # Pega a última linha da tabela
        tbl_element = table._tbl
        tr_element = row_to_delete._tr
        tbl_element.remove(tr_element)
        print(f"INFO: Linha excedente removida. Total de linhas agora: {len(table.rows)}")

    # Adicionar linhas se estiverem faltando
    while len(table.rows) < target_total_rows_in_table:
        new_row = table.add_row()
        new_row.height = Cm(0.5) # Aplicar altura padrão às novas linhas
        new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Garante que as células da nova linha tenham parágrafos formatados
        for cell in new_row.cells:
            # Assegura que existe pelo menos um parágrafo e o alinha
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Pode-se definir a fonte padrão para o parágrafo ou para um run vazio aqui, se necessário
            # Mas geralmente o estilo da tabela ou a formatação do conteúdo adicionado depois cuidará disso.
        print(f"INFO: Linha faltante adicionada. Total de linhas agora: {len(table.rows)}")

    # 3. Preencher as linhas de dados (seu código original a partir daqui)
    for i in range(quantidade_dias_no_mes):
        dia = i + 1
        # Agora é seguro acessar table.rows[linha_inicial + i]
        row = table.rows[linha_inicial + i]
        data_atual = date(ano, mes_numerico, dia) # Use o nome data_atual como no seu código
        dia_semana = pega_final_de_semana(ano, mes_numerico, dia) # Assume que esta função existe

        # Limpeza das células da linha atual antes de preencher
        for cell in row.cells:
            cell.text = "" # Limpa o conteúdo principal da célula (primeiro parágrafo)
            for paragraph in cell.paragraphs: # Itera sobre todos os parágrafos
                paragraph.clear() # Limpa todos os 'runs' (texto formatado) de cada parágrafo
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Garante o alinhamento

        # Preencher número do dia
        dia_cell = row.cells[0]
        # Garante que há um parágrafo para adicionar o run
        dia_paragraph = dia_cell.paragraphs[0] if dia_cell.paragraphs else dia_cell.add_paragraph()
        dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Reafirma o alinhamento se for um novo parágrafo
        dia_run = dia_paragraph.add_run(str(dia))
        dia_run.font.name = "Calibri"
        dia_run.font.size = Pt(8)
        # dia_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Já definido

        # Lógica para Sábados, Domingos, Feriados e Férias
        # A ordem das suas verificações originais define a prioridade (a última condição que sobrescreve cell.text vence)
        
        texto_status = "" # Para Sábado/Domingo

        if dia_semana == 5:
            texto_status = "SÁBADO"
        elif dia_semana == 6:
            texto_status = "DOMINGO"

        if texto_status: # Escreve SÁBADO ou DOMINGO
            for j in [2, 5, 9, 13]: # Seus índices de coluna originais
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = texto_status # Define o texto, limpando parágrafos anteriores
                    # Reaplicar formatação após cell.text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs: # O texto agora está em um ou mais runs
                            run.font.bold = True
                            run.font.name = "Calibri" # Garantir consistência
                            run.font.size = Pt(7)     # Garantir consistência
                else:
                    print(f"AVISO: Índice de coluna {j} para S/D fora dos limites.")


        # Feriado (exceto se for sábado ou domingo) - sobrescreve células se for o caso
        if data_atual in feriados and dia_semana not in [5, 6]:
            for j in [2, 5, 9, 13]:
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = "FERIADO"
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(7)
                else:
                    print(f"AVISO: Índice de coluna {j} para FERIADO fora dos limites.")


        # Férias (exceto fins de semana) - sobrescreve células se for o caso
        if funcionario.get('feriasinicio') and funcionario.get('feriasfinal'):
            ferias_inicio_raw = funcionario['feriasinicio']
            ferias_final_raw = funcionario['feriasfinal']
            ferias_inicio = ferias_inicio_raw.date() if hasattr(ferias_inicio_raw, 'date') else ferias_inicio_raw
            ferias_final = ferias_final_raw.date() if hasattr(ferias_final_raw, 'date') else ferias_final_raw

            if isinstance(ferias_inicio, date) and isinstance(ferias_final, date) and \
               (ferias_inicio <= data_atual <= ferias_final and dia_semana not in [5, 6]):
                for j in [2, 5, 9, 13]:
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = "FÉRIAS"
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = "Calibri"
                                run.font.size = Pt(7)
                    else:
                        print(f"AVISO: Índice de coluna {j} para FÉRIAS fora dos limites.")
    
    # Se você tem múltiplas tabelas no documento e só quer processar a primeira,
    # o loop `for table in doc.tables:` pode ser removido ou adicionar um `break` no final.
    # Se o código original processava todas, mantenha o loop. Pela sua estrutura, parece que
    # a intenção é processar a primeira tabela de frequência encontrada.
    # Adicionando um break para processar apenas a primeira tabela, que é o comportamento mais comum.
    # Remova este 'break' se você intencionalmente processa múltiplas tabelas de frequência no mesmo doc.
    # No seu código original, não havia break, então o loop for table continuaria.
    # Para este ajuste de linhas, faz mais sentido focar em UMA tabela principal.
    # Se o loop for table for mantido, a lógica de ajuste de linhas será aplicada a cada tabela.
